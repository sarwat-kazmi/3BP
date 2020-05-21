import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import sys
from argparse import ArgumentParser
from docx import Document
from docx.shared import Inches

""" Classes for calculation and visual display of finances"""

class Calculation:
    """ Calculating and building various components of a user's financial 
    profile.
    
    Attributes:
        total_income (float): users total income (yearly)
        fixed_expenses (dict):  contains name and dollar amount of 
                                fixed expenses (monthly)
        var_expenses (dict):  contains name and dollar amount of 
                              variable expenses (monthly)
    
    Side effects:
        creates a CSV file 'expense_sheet.csv',
        creates a Word document, 'finances.docx'
            
    Returns:
        CSV file       
    """
    
    def __init__(self, total_income):
        self.total_income = total_income
        self.fixed_expenses = {}
        self.var_expenses = {}
        
    def user_expenses(self):
        """ Asks user to enter fixed and variable expenses and stores that
        information in a DataFrame
        
        Side effects:
            modifies self.fixed_expenses, self.var_expenses,
            creates 'expense_sheet.csv' file,
            asks user for input, 
            writes return result to 'finances.docx'

        Returns:
            (DataFrame):  contains fixed and variable expenses
        
        """
        
        income_df = pd.DataFrame()
        
        while True:
            try:
                fixed = input("FIXED expenses, name? (one word) ")
                if fixed == 'done':
                    break
                if fixed.isalpha() or fixed == " ":
                    pass
                else:
                    raise TypeError
            except TypeError:
                print("Letters only please.")
                continue
            try:
                cost = input("How much?  (dollars and cents) ")
                if cost == 'done':
                    break
                cost = round(float(cost), 2)
            except ValueError:
                print("Numbers only please.")
                continue
            self.fixed_expenses[fixed] = cost

        while True:
            try:
                var = input("VARIABLE expenses, name? ")
                if var == "done":
                    break
                if var.isalpha():
                    pass
                else:
                    raise TypeError
            except TypeError:
                print("Letters only please.")
                continue
            try:
                cost = input("How much?  (dollars and cents) ")
                if cost == 'done':
                    break
                cost = round(float(cost), 2)
            except ValueError:
                print("Numbers only please.")
                continue
            self.var_expenses[var] = cost

        if len(self.fixed_expenses) == 0 and len(self.var_expenses) == 0:
            print("Please enter FIXED and VARIABLE expenses to continue,", 
                  "goodbye.")
            exit()
        
        income_df = pd.DataFrame({'fixed expenses':self.fixed_expenses, 
                                  'variable expenses':self.var_expenses})

        income_df.fillna(0).to_csv('expense_sheet.csv', encoding='utf-8', 
                                   index=False)
        return income_df.fillna(0)
        
    def spend_check(self):
        """ Calculates whether user is overspending by adding fixed and variable
        epxenses and comparing to the users total income.

        Side effects:
            writes return result to 'finances.docx'

        Returns:
            (str):  lets user know if they are over spending or on budget.
        """

        if len(self.fixed_expenses) == 0 & len(self.var_expenses) == 0:
            return ('Please enter both your fixed and variable expenses to run' 
                    + ' this program.')
        
        spend = ""
        fixed = sum(self.fixed_expenses.values())
        var = sum(self.var_expenses.values())

        overspend = round(fixed + var, 2)
        monthly_earn = round(self.total_income/12, 2)
        excess_spend = round(overspend - monthly_earn, 2)

        if fixed + var > monthly_earn:
            spend = ('Every month you are spending $' + str(overspend) + '.' +
                    '  This is $' + str(excess_spend) + 
                    ' over your monthly earnings of $' + 
                    str(monthly_earn))
        else:
            spend = ('Every month you are spending $' + str(overspend) + '.' +
                    '  This is under your monthly earnings of $' + 
                    str(monthly_earn))
        return spend

    def input_info(self):
        """  Asks the user to enter loan information for the interest() method.
        
        Side effects:
            asks the user to enter input,
            prints to the terminal
        
        Returns:
            (dictionary):  contains the input provided by the user
        """
        
        info = {}
        
        while True:
            skip = input("Enter any key to proceed or type 'skip'" + 
                         " to skip this step: ")
            if skip == 'skip':
                return info
            else:
                pass   
            try:
                n = input("Enter NAME of loan (one word): ")
                if n.isalpha():
                    info[n] = 0
                    pass
                else:
                    raise TypeError
            except TypeError:
                print("Letters only please.")
                continue
            try:
                p = input("Enter the PRINCIPAL (dollars and cents): ")
                p = float(p)
                info[p] = 0
            except ValueError:
                print("Numbers only please.")
                continue
            try:
                r = input("Enter the RATE (% per year): ")
                r = float(r)
                info[r] = 0
            except ValueError:
                print("Numbers only please.")
                continue
            try:
                t = input("Enter the TIME (months): ")
                t = int(t)
                info[t] = 0
            except ValueError:
                print("Numbers only please.")
                continue
            try:
                int_type = input("Enter SIMPLE or COMPOUND: ")
                if int_type == 'SIMPLE':
                    info[int_type] = 0
                    return info
                elif int_type == 'COMPOUND':
                    info[int_type] = 0
                    try:
                        inp = input("Enter MONTHLY, QUARTERLY, or YEARLY: ")
                        if (inp == "MONTHLY" or inp == "QUARTERLY" or 
                            inp == "YEARLY"): 
                            info[inp] = 0
                            break
                        else:
                            raise TypeError
                    except TypeError:
                        print("Letters only please.")
                        continue
                else:
                    raise TypeError
            except TypeError:
                print("Enter SIMPLE or COMPOUND please.")
                continue

        return info

    def interest(self): 
        """Calculates the amount of interest that will be accumulated on loans.
        
        Side effects:
            prints to the terminal,
            writes the return result to 'finances.docx'

        Returns:
            (float): the total amount of interest due at the end of the
            loan period.
        """
        lst = []

        for key, val in self.input_info().items():
            lst.append(key)

        if len(lst) == 0:
            interest_due = "No information entered."
            return interest_due
        
        R = lst[2]/100
        T = lst[3]/12
        
        if "SIMPLE" in lst:
            interest_due = ('Loan: ' + str(lst[0]) + ', ' + 
                                 'Principal amount: $' + str(lst[1]) + 
                                 ', ' + 'Rate: ' + str(lst[2]) + '%' + 
                                 ', ' + 'Time: ' + str(lst[3]) + 
                                 ' months' + ', ' +
                                 'Simple Interest Loan' + ', ' + 
                                 'Total interest paid is $' + 
                                 str(round((lst[1] * R * T), 2)))
            
        if "COMPOUND" and "MONTHLY" in lst:
            total_due = lst[1] * (1 + R/12)**(12*T)
            interest_due = ('Loan: ' + str(lst[0]) + ', ' + 
                            'Principal amount: $' + str(lst[1]) + 
                            ', ' + 'Rate: ' + str(lst[2]) + '%' + 
                            ', ' + 'Time: ' + str(lst[3]) + 
                            ' months' + ', ' +
                            'Compounded: Monthly' + ', ' + 
                            'Total interest paid is $' + 
                            str(round(total_due - lst[1], 2)))
        
        if "COMPOUND" and "QUARTERLY" in lst:
            total_due = lst[1] * (1 + R/4)**(4*T)
            interest_due = ('Loan: ' + str(lst[0]) + ', ' + 
                            'Principal amount: $' + str(lst[1]) + 
                            ', ' + 'Rate: ' + str(lst[2]) + '%' + 
                            ', ' + 'Time: ' + str(lst[3]) + 
                            ' months' + ', ' +
                            'Compounded: Quarterly' + ', ' + 
                            'Total interest paid is $' + 
                            str(round(total_due - lst[1], 2)))
            
        if "COMPOUND" and "YEARLY" in lst:
            total_due = lst[1] * (1 + R/1)**(1*T)
            interest_due = ('Loan: ' + str(lst[0]) + ', ' + 
                            'Principal amount: $' + str(lst[1]) + 
                            ', ' + 'Rate: ' + str(lst[2]) + '%' + 
                            ', ' + 'Time: ' + str(lst[3]) + 
                            ' months' + ', ' +
                            'Compounded: Yearly' + ', ' + 
                            'Total interest paid is $' + 
                            str(round(total_due - lst[1], 2)))

        print(interest_due)
        return interest_due

    def spending_allocation(self): 
        """Calculates the percentage of income that goes towards fixed/variable 
        expenses.
        
        Side effects:
            prints to the terminal,
            writes the return result to 'finances.docx'
    
        Returns:
            (tuple): the percentage of income going towards each category.
        """
        if len(self.fixed_expenses) == 0 & len(self.var_expenses) == 0:
            return ('Please enter both your fixed and variable expenses to run' 
                    + ' this program.')
        
        fixed = sum(self.fixed_expenses.values())
        var = sum(self.var_expenses.values())

        pctg1 = round((fixed/self.total_income)*100, 2)
        pctg2 = round((var/self.total_income)*100, 2)
        return pctg1, pctg2

class Graphs:
    """  A visual representation of user spending
    
    Attributes:
        expense (dict):  name and dollar amount of users monthly expenses

    Returns:
        pie chart, bar plot
    """
    
    def __init__(self):
        self.expense = {}

    def expenses(self, fixed, var):
        """ Provides a visual representation of the users expenses
        
        Args:
            fixed (dict):  name and dollar amount of fixed expenses (monthly)
            var (dict):  name and dollar amount of variable expenses (monthly)
        
        Side effects:
            modifies self.expense, 
            prints to the terminal, 
            displays visual graphic

        Returns:
            (bar plot):  displays users' expenses
        """
        if len(fixed) == 0 & len(var) == 0:
            return ('Please enter both your fixed and variable expenses to run' 
                    + ' this program.')

        for x in [fixed, var]:
            self.expense.update(x)
        keys = self.expense.keys()
        values = self.expense.values()
        
        fig, ax = plt.subplots()
        bp = ax.bar(keys, values)
        
        ax.set_ylabel("Amount in dollars")
        ax.set_title("User Fixed and Variable Expenses")
        ax.set_xlabel("Expense Name")
        return plt.show()
    
    def input_habits(self):
        """  Asks the user to enter a list of their habits and their monthly 
        costs
        
        Side effects:
            prints to the terminal,
            requests user input
        
        Returns:
            (dictionary):  contains the users' listed habits with corresponding 
            dollar amount
        """
        
        habits = {}

        while True:
            try:
                habit = input("Enter NAME of habit (one word): ")
                if habit == "done":
                    break
                if habit == "none":
                    break
                if habit.isalpha():
                    pass
                else:
                    raise TypeError
            except TypeError:
                print("Letters only please.")
                continue
            try:
                cost = input("How much do you spend monthly to supplement this " +
                                "habit? (enter dollars and cents): ")
                cost = float(cost)
                habits[habit] = cost
                continue
            except ValueError:
                print("Numbers only please.")
                continue
    
        return habits
        
    def habits(self, income):
        """  Provides a visual representation of how much money the user 
        is spending on their listed habits
        
        Args:
            income (float):  users total income (yearly amount)
        
        Side effects:
            prints to the terminal, 
            displays visual graphic,
            writes to 'finances' document
        
        Returns:
            (pie chart):  displays users' spending habits  
        """
        
        hab = [] 
        total = []
        saving = 0.0
        lst = []

        for k, v in self.input_habits().items():
            hab.append(k)
            total.append(v)
            print('You are spending ' + str(v) + 
                  ' dollar(s) every month on ' + '[' + k + ']' + '.')
            print('This means every year you are spending ' + 
                  str(round(v*12, 2)) + ' dollar(s) on ' + '[' + k + ']' + '.')
            print('[' + k + ']' + ' account(s) for ' + 
                  str(round(((v*12)/income)*100, 2)) + 
                  ' percent of your total income per year.\n')
            saving += v*12
            lst.append((k, v, v*12))
        
        print('You are spending $' + str(round(saving, 2)) + 
              ' dollar(s) a year on your habits.')
        
        for_doc = (lst)
        if len(for_doc) == 0:
            for_doc = "No habits entered."
        
        labels = hab
        sizes = total
        
        fig, ax = plt.subplots()
        ax.pie(sizes, labels=labels, autopct = '%1.1f%%', startangle=90)
        ax.axis('equal')
        ax.set_title("Habits")
        
        return for_doc, plt.show()

def parse_args(arglist):
    parser = ArgumentParser()
    parser.add_argument("total_income", help="total income for current year",
                        type=float)
    args = parser.parse_args(arglist)
    return args

def main(arglist):
    count = 0
    lst = set()
    lst2 = set()
    
    args = parse_args(arglist)
    c = Calculation(args.total_income)
    g = Graphs()
    
    d = Document()
    
    d.add_heading("Three Blind Programmers Presents:  A Birds-eye-view" 
                  " of your finances", 0)
    d.add_paragraph('Yearly income: $' + str(c.total_income))
    d.add_heading('Monthly fixed and variable expenses.', 1)
    
    print(c.user_expenses())
    for exp, num in c.fixed_expenses.items():
        count += 1
        lst.add(str((exp, num)))
    
    count = 0
    for exp, num in c.var_expenses.items():
        count += 1
        lst2.add(str((exp, num)))
    
    table = d.add_table(rows=count, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Fixed Expenses'
    hdr_cells[1].text = 'Variable Expenses'

    for num in lst:
        row_cells = table.add_row().cells
        row_cells[0].text = num
    
    for num in lst2:
        row_cells = table.add_row().cells
        row_cells[1].text = num
    
    print(c.spend_check())
    d.add_heading('Spending per month.', 1)
    d.add_paragraph(str(c.spend_check()))
    
    d.add_heading('Loan information.', 1)
    d.add_paragraph(str(c.interest()))

    print('Fixed/variable spending, respectively (percent of total income): ' 
          + str(c.spending_allocation()))
    d.add_heading('Monthly expenses against total income: $' + 
                  str(c.total_income), 1)
    d.add_paragraph('Fixed expenses account for ' + 
                    str(c.spending_allocation()[0]) + '%' + 
                    ' of your total income, totaling $' + 
                    str(sum(c.fixed_expenses.values())) + ' a month.')
    d.add_paragraph('Variable expenses account for ' + 
                    str(c.spending_allocation()[1]) + '%' + 
                    ' of your total income, totaling $' + 
                    str(sum(c.var_expenses.values())) + ' a month.')
    
    print(g.expenses(c.fixed_expenses, c.var_expenses))

    d.add_heading('Monthly habit expenses.', 1)
    d.add_paragraph('Habit, Monthly cost, Yearly cost')
    d.add_paragraph(str(g.habits(c.total_income)[0]))
    
    d.save('finances.docx')
  
if __name__ == "__main__":
    main(sys.argv[1:])